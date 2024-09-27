"""
@file: 文档内容补充ABCD.py
@time: 2024/6/13 15:05
@desc: 
"""


head_content = """您与调研对象的关系是<[input]>

调研对象的姓名<[input]>

调研对象的年龄<[input]>

调研对象的职业<[input]>

调研对象的婚姻状况
未婚
已婚

调研对象所在城市<[input]>"""

directory = "/workspace/python_learn_record/baobao/文档内容补充ABCD/问卷0612"

# 将 head_content 的内容写入到 directory 路径下的每一个 docx 文件的开头,
# 其中 路径 directory 是一个包含多个 docx 文件的目录
# 文档中有很多 A. 经常  B.  偶尔  C.  很少  D.  从未   这样的选项，需要调整为 竖排的形式
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# 遍历指定目录下的所有docx文件
def get_all_docx_file(directory):
    docx_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(".docx"):
                docx_files.append(os.path.join(root, file))
    return docx_files


# 将 head_content 的内容写入到 docx 文件的开头
def write_head_content_to_docx(docx_file):
    """
    1、将 head_content 的内容写入到 docx 文件的开头
    2、文档中有很多 A. 经常  B.  偶尔  C.  很少  D.  从未   这样的选项，需要调整为 竖排的形式
    :param docx_file:
    :return:
    """
    doc = Document(docx_file)
    # 获取文档的第一段
    new_doc = Document()
    # 添加头部内容到新文档
    # new_doc.add_paragraph(head_content)
    p = new_doc.add_paragraph()
    run = p.add_run(head_content)
    set_font_to_song(run)

    # 文档中有很多 A. 经常  B.  偶尔  C.  很少  D.  从未   这样的选项，需要调整为 竖排的形式
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "A" not in run.text:
                continue
            # for i in run.text.split("\n\n"):
            #     for j in i.split("\n"):
            run.text = run.text.replace(" B.", "\nB.").replace(" C.", "\nC.").replace(" D.", "\nD.").replace(" E.", "\nE.").replace(" F.", "\nF.").replace(" G.", "\nG.")
            print(run.text)
            set_font_to_song(run)

    # 将原文档的内容添加到新文档, 使用宋体
    for element in doc.element.body:
        new_doc.element.body.append(element)

    new_doc.save(docx_file)


def set_font_to_song(run):
    run.font.name = '宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def write_head_content_to_docx1(docx_file):
    doc = Document(docx_file)
    # 在文档开头插入内容
    # doc.insert_paragraph_before(head_content)
    new_doc = Document()
    # 添加头部内容到新文档
    new_doc.add_paragraph(head_content)

    # 将原文档的内容添加到新文档
    for element in doc.element.body:
        new_doc.element.body.append(element)

    # 保存新文档覆盖原文件
    new_doc.save(docx_file)


if __name__ == '__main__':
    docx_files = get_all_docx_file(directory)
    for docx_file in docx_files:
        # write_head_content_to_docx(docx_file)
        write_head_content_to_docx(docx_file)

    print("所有文件已成功更新。")