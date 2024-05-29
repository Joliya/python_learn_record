"""
@file: excel_to_word.py
@time: 2023/12/19 16:26
@desc: 
"""
import os
import sys

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


def get_executable_dir():
    if getattr(sys, 'frozen', False):
        # 脚本被冻结，获取可执行文件的路径
        home_directory = os.path.dirname(sys.executable)
    else:
        # 脚本未被冻结，获取脚本所在的路径
        home_directory = os.path.dirname(os.path.abspath(__file__))
    return home_directory


if __name__ == '__main__':
    home_directory = get_executable_dir()
    data = pd.read_excel(f"{home_directory}/27篇乳腺癌.xlsx")

    for i in data.values:
        disease, clinic, title, content = i
        path = f"{home_directory}/{clinic}/{disease}"
        if not os.path.exists(path):
            os.makedirs(path)
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = u'SimSun'
        font._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        document.add_heading(title, 0)
        paragraph = document.add_paragraph(content)
        run = paragraph.runs[0]
        run.font.size = Pt(11)
        document.save(f"{path}/{title}.docx")
